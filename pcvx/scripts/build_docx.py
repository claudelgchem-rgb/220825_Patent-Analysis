#!/usr/bin/env python3
"""1번 산출물 — 워드 보고서 docx.

a9_report.md(본문)를 뼈대로 삼고, 표지·경고문·한계 원문·부록을 데이터에서 직접 만든다.
등급 색상: 상=진한 녹색, 중=주황, 하=빨강.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
import common  # noqa: E402

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

GRADE_RE = re.compile(r"\[(상|중|하)\]")


def rgb(hexcode: str) -> RGBColor:
    return RGBColor(int(hexcode[0:2], 16), int(hexcode[2:4], 16), int(hexcode[4:6], 16))


def add_graded_text(paragraph, text: str) -> None:
    """[상]/[중]/[하] 표시에 색을 입혀 문단을 만든다."""
    pos = 0
    for match in GRADE_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        run = paragraph.add_run(match.group(0))
        run.bold = True
        run.font.color.rgb = rgb(common.GRADE_COLORS[match.group(1)])
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_cover(doc: Document, env: dict, topic: dict, attempts: dict) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("특허 조사 보고서")
    run.font.size = Pt(30)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run(common.cell(topic.get("topic"), "조사 주제 미확정"))
    srun.font.size = Pt(16)

    doc.add_paragraph()
    meta = doc.add_table(rows=0, cols=2)
    meta.style = "Light Grid Accent 1"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("작성 기준일", env["as_of"]),
        ("조사 대상 관할국", common.cell(topic.get("jurisdictions"))),
        ("조사 대상 기간", common.cell(topic.get("period"))),
        ("출원인 한정", common.cell(topic.get("assignee_filter"), "한정 없음")),
        ("목표 건수", common.cell(topic.get("target_count"))),
    ]
    sources = [a.get("source") for a in attempts.get("attempts", []) if a.get("attempted")]
    rows.append(("검색 대상 데이터베이스", common.cell(sorted(set(sources)))))
    for label, value in rows:
        cells = meta.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
    doc.add_page_break()


def add_warning(doc: Document, low_ratio: float) -> None:
    if low_ratio <= 0.30:
        return
    para = doc.add_paragraph()
    run = para.add_run(common.low_ratio_warning())
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = rgb(common.GRADE_COLORS["하"])
    doc.add_paragraph(
        f"신뢰도 '하' 등급 항목이 전체의 {low_ratio * 100:.1f}퍼센트를 차지한다."
    )
    doc.add_page_break()


def md_table_to_doc(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for idx, head in enumerate(rows[0]):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        add_graded_text(cell.paragraphs[0], head)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx in range(len(rows[0])):
            value = row[idx] if idx < len(row) else common.FAILURE_MARK
            cells[idx].text = ""
            add_graded_text(cells[idx].paragraphs[0], common.cell(value))


def render_markdown(doc: Document, markdown: str) -> None:
    """A9 본문 마크다운을 워드 요소로 옮긴다. 표는 표로, 제목은 제목으로."""
    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if stripped.startswith("|") and stripped.endswith("|"):
            block = []
            while idx < len(lines):
                current = lines[idx].strip()
                if not (current.startswith("|") and current.endswith("|")):
                    break
                cells = [c.strip() for c in current.strip("|").split("|")]
                if not all(set(c) <= set("-: ") for c in cells):
                    block.append(cells)
                idx += 1
            md_table_to_doc(doc, block)
            doc.add_paragraph()
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            doc.add_heading(stripped.lstrip("# ").strip(), level=min(level, 4))
        elif stripped.startswith(("- ", "* ")):
            para = doc.add_paragraph(style="List Bullet")
            add_graded_text(para, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            para = doc.add_paragraph(style="List Number")
            add_graded_text(para, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith(">"):
            para = doc.add_paragraph()
            run = para.add_run(stripped.lstrip("> ").strip())
            run.italic = True
        elif stripped:
            para = doc.add_paragraph()
            add_graded_text(para, stripped)
        else:
            doc.add_paragraph()
        idx += 1


def add_limits(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("7. 이 조사의 한계", level=1)
    for block in common.limits_text().split("\n\n"):
        doc.add_paragraph(block.strip())


def add_appendix(doc: Document, queries: dict, attempts: dict, records: list) -> None:
    doc.add_page_break()
    doc.add_heading("8. 부록", level=1)

    doc.add_heading("8.1 검색식 전문", level=2)
    doc.add_paragraph(
        "아래 표는 이번 조사에서 실제로 사용한 검색식을 모두 옮긴 것이다. "
        "왼쪽부터 식 번호, 유형, 검색식 본문, 그 식으로 무엇을 잡으려 했는지의 순서로 읽으면 된다. "
        "같은 주제라도 검색식이 다르면 결과가 달라지므로, 재현이 필요할 때 이 표를 그대로 쓰면 된다."
    )
    rows = [["식 번호", "유형", "검색식", "의도"]]
    for q in queries.get("queries", []):
        rows.append([common.cell(q.get("id")), common.cell(q.get("type")),
                     common.cell(q.get("expression")), common.cell(q.get("intent"))])
    md_table_to_doc(doc, rows)

    doc.add_heading("8.2 출처 목록과 접속일", level=2)
    doc.add_paragraph(
        "아래 표는 어떤 데이터베이스를 언제 열어 무엇을 얻었는지 기록한 것이다. "
        "접속일이 다르면 법적 상태가 달라져 있을 수 있으므로, 날짜를 함께 보는 것이 중요하다. "
        "접근이 막혀 실패한 시도도 숨기지 않고 그대로 남겨 두었다."
    )
    rows = [["소스", "URL", "시간 창", "수집 건수", "접속일", "비고"]]
    today = common.today()
    for a in attempts.get("attempts", []):
        rows.append([
            common.cell(a.get("source")), common.cell(a.get("url")),
            common.cell(a.get("window")), common.cell(a.get("hits"), "0"),
            common.cell(a.get("accessed_on"), today),
            common.cell(a.get("failure") or a.get("note"), "정상 수행"),
        ])
    for rec in records:
        rows.append([
            common.cell(rec.get("representative_doc")), common.cell(rec.get("source_url")),
            "해당 없음", "1", common.cell(rec.get("legal_status_checked_on"), today), "문헌 원문",
        ])
    md_table_to_doc(doc, rows)


def main() -> int:
    env = common.env()
    topic = common.read_json(common.WORKSPACE / "topic.json", {}) or {}
    body_path = common.WORKSPACE / "a9_report.md"
    if not body_path.exists():
        print("[실패] a9_report.md 가 없다. A9 를 먼저 실행한다.")
        return 1
    records = (common.read_json(common.WORKSPACE / "records.json", {}) or {}).get("records", [])
    verification = common.read_json(common.WORKSPACE / "a7_verification.json", {}) or {}
    queries = common.read_json(common.WORKSPACE / "a1_queries.json", {}) or {}
    attempts = common.read_json(common.WORKSPACE / "source_attempts.json", {}) or {}

    dist = verification.get("distribution", {})
    total = sum(dist.get(g, 0) for g in common.GRADES) or 1
    low_ratio = dist.get("low_ratio") or (dist.get("하", 0) / total)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(10.5)

    add_cover(doc, env, topic, attempts)
    add_warning(doc, float(low_ratio))
    render_markdown(doc, body_path.read_text(encoding="utf-8"))
    add_limits(doc)
    add_appendix(doc, queries, attempts, records)

    slug = common.slugify(topic.get("topic", "무제"))
    out = common.output_dir() / f"PCVX_특허조사보고서_{slug}_{env['as_of_compact']}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"[기준일: {env['as_of']} / 도구: build_docx]")
    print(f"-> {out}  ({out.stat().st_size:,} bytes, 문단 {len(doc.paragraphs)}개, 표 {len(doc.tables)}개)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
